"""
FutureClarity Automation - AI Chatbot System
A complete RAG-based chatbot system using free, open-source tools.

Features:
- Local LLM integration via Ollama
- Document-based RAG with Chroma vector database
- FastAPI backend with RESTful endpoints
- Embeddable web interface
- Easy document management

Requirements:
- Python 3.8+
- Ollama installed locally
- Internet connection for initial setup

Author: FutureClarity Automation
License: MIT
"""

import os
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import asyncio
import aiofiles
import hashlib

# Web framework and async support
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel
import uvicorn

# RAG and ML libraries
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
from markdown import markdown
from bs4 import BeautifulSoup
import requests

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Custom middleware to add ngrok bypass headers for mobile compatibility

# Configuration
CONFIG = {
    "app_name": "FutureClarity Chatbot",
    "version": "1.0.0",
    "ollama_base_url": "http://localhost:11434",
    "default_model": "llama3.1:8b",
    "embedding_model": "all-MiniLM-L6-v2",
    "chroma_db_path": "./chroma_db",
    "documents_path": "./documents",
    "max_context_length": 4000,
    "chunk_size": 500,
    "chunk_overlap": 50,
    "top_k_results": 3,
    "port": 3000,
    "host": "0.0.0.0"
}

# Pydantic models
class ChatMessage(BaseModel):
    message: str
    conversation_id: Optional[str] = None

class ChatResponse(BaseModel):
    response: str
    conversation_id: str
    sources: List[str] = []

class DocumentInfo(BaseModel):
    filename: str
    size: int
    processed_at: str
    chunks: int

# FastAPI app will be initialized after lifespan handler is defined

class ChatbotSystem:
    def __init__(self):
        self.embedding_model = None
        self.chroma_client = None
        self.collection = None
        self.documents_processed = {}
        self.conversation_history = {}
        
    async def initialize(self):
        """Initialize the chatbot system components"""
        logger.info("Initializing chatbot system...")
        
        # Create directories
        os.makedirs(CONFIG["documents_path"], exist_ok=True)
        os.makedirs(CONFIG["chroma_db_path"], exist_ok=True)
        
        # Initialize embedding model
        logger.info("Loading embedding model...")
        self.embedding_model = SentenceTransformer(CONFIG["embedding_model"])
        
        # Initialize Chroma database
        logger.info("Setting up vector database...")
        try:
            self.chroma_client = chromadb.PersistentClient(
                path=CONFIG["chroma_db_path"],
                settings=Settings(anonymized_telemetry=False, allow_reset=True)
            )
        except Exception as e:
            logger.error(f"Error initializing Chroma client: {e}")
            raise
        
        # Get or create collection
        try:
            self.collection = self.chroma_client.get_collection("documents")
            logger.info("Loaded existing document collection")
        except Exception as e:
            logger.info(f"Collection not found, creating new one: {e}")
            try:
                self.collection = self.chroma_client.create_collection(
                    name="documents",
                    metadata={"hnsw:space": "cosine"}
                )
                logger.info("Created new document collection")
            except Exception as create_error:
                logger.error(f"Error creating collection: {create_error}")
                self.chroma_client.reset()
                self.collection = self.chroma_client.create_collection(
                    name="documents",
                    metadata={"hnsw:space": "cosine"}
                )
                logger.info("Reset database and created new collection")
        except Exception as e:
            logger.error(f"Error initializing ChromaDB: {e}")
            # Create fresh database
            logger.info("Creating fresh ChromaDB instance...")
            self.chroma_client = chromadb.PersistentClient(
                path=CONFIG["chroma_db_path"],
                settings=Settings(anonymized_telemetry=False, allow_reset=True)
            )
            self.collection = self.chroma_client.create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"}
            )
            logger.info("Created fresh document collection")
        
        # Load processed documents metadata
        self.load_processed_documents()
        
        # Check Ollama connection
        await self.check_ollama_connection()
        
        logger.info("Chatbot system initialized successfully!")
    
    def load_processed_documents(self):
        """Load metadata about processed documents"""
        metadata_file = Path(CONFIG["documents_path"]) / "processed_documents.json"
        if metadata_file.exists():
            with open(metadata_file, 'r') as f:
                self.documents_processed = json.load(f)
    
    def save_processed_documents(self):
        """Save metadata about processed documents"""
        metadata_file = Path(CONFIG["documents_path"]) / "processed_documents.json"
        with open(metadata_file, 'w') as f:
            json.dump(self.documents_processed, f, indent=2)
    
    async def check_ollama_connection(self):
        """Check if Ollama is running and model is available"""
        try:
            response = requests.get(f"{CONFIG['ollama_base_url']}/api/tags")
            if response.status_code == 200:
                models = response.json().get("models", [])
                model_names = [model["name"] for model in models]
                
                if CONFIG["default_model"] in model_names:
                    logger.info(f"Ollama model '{CONFIG['default_model']}' is available")
                else:
                    logger.warning(f"Model '{CONFIG['default_model']}' not found. Available models: {model_names}")
                    if model_names:
                        CONFIG["default_model"] = model_names[0]
                        logger.info(f"Using model: {CONFIG['default_model']}")
                    else:
                        logger.error("No models available in Ollama")
            else:
                logger.error("Failed to connect to Ollama")
        except Exception as e:
            logger.error(f"Error checking Ollama: {str(e)}")
    
    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_path.suffix.lower() == '.md':
                return self.extract_text_from_markdown(file_path)
            elif file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logger.warning(f"Unsupported file format: {file_path.suffix}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def extract_text_from_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert markdown to HTML then extract text
        html = markdown(markdown_content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks"""
        if chunk_size is None:
            chunk_size = CONFIG["chunk_size"]
        if overlap is None:
            overlap = CONFIG["chunk_overlap"]
        
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks
    
    async def process_document(self, file_path: Path) -> DocumentInfo:
        """Process a document and add it to the vector database"""
        logger.info(f"Processing document: {file_path.name}")
        
        # Extract text
        text = self.extract_text_from_file(file_path)
        if not text.strip():
            raise ValueError("No text could be extracted from the document")
        
        # Create chunks
        chunks = self.chunk_text(text)
        
        # Generate embeddings
        embeddings = self.embedding_model.encode(chunks).tolist()
        
        # Create unique IDs for chunks
        file_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
        chunk_ids = [f"{file_hash}_{i}" for i in range(len(chunks))]
        
        # Prepare metadata
        metadata = [
            {
                "filename": file_path.name,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "processed_at": datetime.now().isoformat()
            }
            for i in range(len(chunks))
        ]
        
        # Add to Chroma collection
        self.collection.add(
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadata,
            ids=chunk_ids
        )
        
        # Update processed documents metadata
        doc_info = DocumentInfo(
            filename=file_path.name,
            size=file_path.stat().st_size,
            processed_at=datetime.now().isoformat(),
            chunks=len(chunks)
        )
        
        self.documents_processed[file_path.name] = doc_info.dict()
        self.save_processed_documents()
        
        logger.info(f"Successfully processed {file_path.name} into {len(chunks)} chunks")
        return doc_info
    
    async def query_documents(self, query: str, top_k: int = None) -> List[Dict[str, Any]]:
        """Query the vector database for relevant documents"""
        if top_k is None:
            top_k = CONFIG["top_k_results"]
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query]).tolist()
        
        # Search in Chroma
        results = self.collection.query(
            query_embeddings=query_embedding,
            n_results=top_k
        )
        
        # Format results
        formatted_results = []
        for i in range(len(results['documents'][0])):
            formatted_results.append({
                'content': results['documents'][0][i],
                'metadata': results['metadatas'][0][i],
                'distance': results['distances'][0][i] if 'distances' in results else 0
            })
        
        return formatted_results
    
    async def generate_response(self, query: str, context: List[str]) -> str:
        """Generate response using Ollama LLM"""
        # Prepare context
        context_text = "\n\n".join(context)
        
        # Create prompt
        prompt = f"""You are a helpful AI assistant for FutureClarity Automation, a company that builds custom AI chatbots, phone voice assistants, and workflow automation for businesses.

Use the following context to answer the user's question. If the answer is not in the context, say "I don't have information about that in my knowledge base."

Context:
{context_text}

Question: {query}

Answer:"""
        
        # Call Ollama API
        try:
            response = requests.post(
                f"{CONFIG['ollama_base_url']}/api/generate",
                json={
                    "model": CONFIG["default_model"],
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.7,
                        "top_p": 0.9,
                        "num_predict": 512
                    }
                },
                timeout=30
            )
            
            if response.status_code == 200:
                return response.json()["response"].strip()
            else:
                logger.error(f"Ollama API error: {response.status_code}")
                return "I'm experiencing technical difficulties. Please try again later."
        
        except Exception as e:
            logger.error(f"Error calling Ollama: {str(e)}")
            return "I'm experiencing technical difficulties. Please try again later."
    
    async def chat(self, message: str, conversation_id: str = None) -> ChatResponse:
        """Process a chat message and return response"""
        if conversation_id is None:
            conversation_id = hashlib.md5(f"{datetime.now().isoformat()}_{message}".encode()).hexdigest()[:16]
        
        # Query relevant documents
        relevant_docs = await self.query_documents(message)
        
        # Extract context and sources
        context = [doc['content'] for doc in relevant_docs]
        sources = list(set([doc['metadata']['filename'] for doc in relevant_docs]))
        
        # Generate response
        response = await self.generate_response(message, context)
        
        # Store conversation history
        if conversation_id not in self.conversation_history:
            self.conversation_history[conversation_id] = []
        
        self.conversation_history[conversation_id].append({
            'timestamp': datetime.now().isoformat(),
            'user_message': message,
            'bot_response': response,
            'sources': sources
        })
        
        return ChatResponse(
            response=response,
            conversation_id=conversation_id,
            sources=sources
        )

# Initialize chatbot system
chatbot = ChatbotSystem()

# Lifespan event handler
from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    await chatbot.initialize()
    yield
    # Shutdown (if needed)
    pass

# Update FastAPI app to use lifespan
app = FastAPI(
    title=CONFIG["app_name"],
    version=CONFIG["version"],
    description="AI Chatbot with RAG capabilities",
    lifespan=lifespan
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(chat_message: ChatMessage):
    """Chat endpoint for processing user messages"""
    try:
        return await chatbot.chat(chat_message.message, chat_message.conversation_id)
    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a document"""
    try:
        # Save uploaded file
        file_path = Path(CONFIG["documents_path"]) / file.filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        # Process document
        doc_info = await chatbot.process_document(file_path)
        
        return JSONResponse(
            content={"message": "Document processed successfully", "info": doc_info.dict()},
            status_code=200
        )
    
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.get("/api/documents")
async def list_documents():
    """List all processed documents"""
    return JSONResponse(content=chatbot.documents_processed)

@app.delete("/api/documents/{filename}")
async def delete_document(filename: str):
    """Delete a processed document"""
    try:
        # Remove from processed documents
        if filename in chatbot.documents_processed:
            del chatbot.documents_processed[filename]
            chatbot.save_processed_documents()
        
        # Remove file
        file_path = Path(CONFIG["documents_path"]) / filename
        if file_path.exists():
            file_path.unlink()
        
        # Remove from vector database (this is simplified - in production you'd want to track and remove specific chunks)
        # For now, we'll just mark it as deleted
        
        return JSONResponse(content={"message": "Document deleted successfully"})
    
    except Exception as e:
        logger.error(f"Delete error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return JSONResponse(content={
        "status": "healthy",
        "version": CONFIG["version"],
        "model": CONFIG["default_model"],
        "documents_processed": len(chatbot.documents_processed)
    })

@app.get("/embed")
async def embed_chatbot():
    """Embeddable chatbot interface"""
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
        
        <!-- Enhanced ngrok bypass headers for mobile -->
        <meta name="ngrok-skip-browser-warning" content="any">
        <meta name="Ngrok-Skip-Browser-Warning" content="any">
        <meta http-equiv="X-Requested-With" content="XMLHttpRequest">
        <meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate">
        <meta http-equiv="Pragma" content="no-cache">
        <meta http-equiv="Expires" content="0">
        
        <!-- Mobile-specific meta tags -->
        <meta name="mobile-web-app-capable" content="yes">
        <meta name="apple-mobile-web-app-capable" content="yes">
        <meta name="apple-mobile-web-app-status-bar-style" content="default">
        <meta name="apple-mobile-web-app-title" content="FutureClarity Chat">
        <meta name="theme-color" content="#667eea">
        
        <title>FutureClarity Chatbot</title>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }
            
            .chatbot-container {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                width: 100%;
                max-width: 400px;
                height: 600px;
                min-height: 400px;
                border: 1px solid #e1e5e9;
                border-radius: 12px;
                background: white;
                display: flex;
                flex-direction: column;
                overflow: hidden;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
            }
            
            /* Mobile responsiveness */
            @media (max-width: 480px) {
                .chatbot-container {
                    max-width: 100%;
                    height: 80vh;
                    min-height: 500px;
                    border-radius: 8px;
                    margin: 0;
                }
                
                .chat-header {
                    padding: 12px;
                    font-size: 16px;
                }
                
                .chat-messages {
                    padding: 12px;
                }
                
                .message-bubble {
                    max-width: 85%;
                    padding: 10px 14px;
                    font-size: 14px;
                    line-height: 1.4;
                }
                
                .chat-input {
                    padding: 12px;
                }
                
                .chat-input input {
                    padding: 12px 16px;
                    font-size: 16px; /* Prevents zoom on iOS */
                    border-radius: 25px;
                }
                
                .chat-input button {
                    padding: 12px 16px;
                    margin-left: 8px;
                    font-size: 14px;
                    min-width: 60px;
                }
            }
            
            .chat-header {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 16px;
                text-align: center;
                font-weight: 600;
            }
            
            .chat-messages {
                flex: 1;
                overflow-y: auto;
                padding: 16px;
                background: #f8f9fa;
            }
            
            .message {
                margin-bottom: 16px;
                display: flex;
                align-items: flex-start;
            }
            
            .message.user {
                justify-content: flex-end;
            }
            
            .message-bubble {
                max-width: 80%;
                padding: 12px 16px;
                border-radius: 18px;
                word-wrap: break-word;
            }
            
            .message.user .message-bubble {
                background: #007bff;
                color: white;
            }
            
            .message.bot .message-bubble {
                background: white;
                color: #333;
                border: 1px solid #e1e5e9;
            }
            
            .message-sources {
                font-size: 0.8em;
                color: #666;
                margin-top: 4px;
                font-style: italic;
            }
            
            .chat-input {
                display: flex;
                padding: 16px;
                background: white;
                border-top: 1px solid #e1e5e9;
            }
            
            .chat-input input {
                flex: 1;
                padding: 12px;
                border: 1px solid #e1e5e9;
                border-radius: 20px;
                outline: none;
                font-size: 14px;
            }
            
            .chat-input button {
                margin-left: 8px;
                padding: 12px 20px;
                background: #007bff;
                color: white;
                border: none;
                border-radius: 20px;
                cursor: pointer;
                font-size: 14px;
                transition: background 0.2s;
            }
            
            .chat-input button:hover {
                background: #0056b3;
            }
            
            .chat-input button:disabled {
                background: #ccc;
                cursor: not-allowed;
            }
            
            .typing-indicator {
                display: none;
                padding: 12px 16px;
                font-style: italic;
                color: #666;
            }
            
            .typing-dots {
                display: inline-block;
                animation: typing 1.5s infinite;
            }
            
            @keyframes typing {
                0%, 60%, 100% { opacity: 1; }
                30% { opacity: 0.5; }
            }
        </style>
    </head>
    <body>
        <div class="chatbot-container">
            <div class="chat-header">
                FutureClarity Assistant
            </div>
            
            <div class="chat-messages" id="chatMessages">
                <div class="message bot">
                    <div class="message-bubble">
                        Hello! I'm your FutureClarity Assistant. I can help you with information about our AI chatbot services, workflow automation, and answer questions based on our knowledge base. How can I assist you today?
                    </div>
                </div>
            </div>
            
            <div class="typing-indicator" id="typingIndicator">
                <span class="typing-dots">Assistant is typing...</span>
            </div>
            
            <div class="chat-input">
                <input type="text" id="messageInput" placeholder="Type your message..." />
                <button id="sendButton">Send</button>
            </div>
        </div>
        
        <script>
            class Chatbot {
                constructor() {
                    this.messagesContainer = document.getElementById('chatMessages');
                    this.messageInput = document.getElementById('messageInput');
                    this.sendButton = document.getElementById('sendButton');
                    this.typingIndicator = document.getElementById('typingIndicator');
                    this.conversationId = null;
                    
                    this.init();
                }
                
                init() {
                    this.sendButton.addEventListener('click', () => this.sendMessage());
                    this.messageInput.addEventListener('keypress', (e) => {
                        if (e.key === 'Enter' && !e.shiftKey) {
                            e.preventDefault();
                            this.sendMessage();
                        }
                    });
                    
                    // Prevent mobile keyboard from resizing viewport
                    this.messageInput.addEventListener('focus', () => {
                        setTimeout(() => {
                            this.messagesContainer.scrollTop = this.messagesContainer.scrollHeight;
                        }, 300);
                    });
                    
                    // Auto-focus on mobile for better UX
                    if (this.isMobile()) {
                        this.messageInput.addEventListener('blur', () => {
                            // Scroll to bottom when keyboard closes
                            setTimeout(() => {
                                this.messagesContainer.scrollTop = this.messagesContainer.scrollHeight;
                            }, 300);
                        });
                    }
                }
                
                isMobile() {
                    return /Android|webOS|iPhone|iPad|iPod|BlackBerry|IEMobile|Opera Mini/i.test(navigator.userAgent);
                }
                
                async sendMessage() {
                    const message = this.messageInput.value.trim();
                    if (!message) return;
                    
                    // Add user message
                    this.addMessage(message, 'user');
                    this.messageInput.value = '';
                    this.setInputDisabled(true);
                    this.showTypingIndicator();
                    
                    try {
                        // Enhanced mobile ngrok bypass headers
                        const headers = {
                            'Content-Type': 'application/json',
                            'ngrok-skip-browser-warning': 'any',
                            'Ngrok-Skip-Browser-Warning': 'any',
                            'X-Requested-With': 'XMLHttpRequest',
                            'Cache-Control': 'no-cache',
                            'Pragma': 'no-cache',
                        };
                        
                        // Add mobile-specific headers
                        if (this.isMobile()) {
                            headers['X-Mobile-Request'] = 'true';
                            headers['X-Requested-With'] = 'XMLHttpRequest';
                            headers['Accept'] = 'application/json, text/plain, */*';
                        }
                        
                        const response = await fetch('/api/chat', {
                            method: 'POST',
                            headers: headers,
                            mode: 'cors',
                            credentials: 'omit',
                            body: JSON.stringify({
                                message: message,
                                conversation_id: this.conversationId
                            })
                        });
                        
                        const data = await response.json();
                        
                        if (response.ok) {
                            this.conversationId = data.conversation_id;
                            this.addMessage(data.response, 'bot', data.sources);
                        } else {
                            this.addMessage('Sorry, I encountered an error. Please try again.', 'bot');
                        }
                    } catch (error) {
                        console.error('Error:', error);
                        // Enhanced error handling for mobile
                        if (this.isMobile() && error.message.includes('network')) {
                            this.addMessage('Network error. Please check your connection and try again.', 'bot');
                        } else {
                        this.addMessage('Sorry, I encountered an error. Please try again.', 'bot');
                        }
                    } finally {
                        this.hideTypingIndicator();
                        this.setInputDisabled(false);
                        this.messageInput.focus();
                    }
                }
                
                addMessage(text, sender, sources = []) {
                    const messageDiv = document.createElement('div');
                    messageDiv.className = `message ${sender}`;
                    
                    const bubbleDiv = document.createElement('div');
                    bubbleDiv.className = 'message-bubble';
                    bubbleDiv.textContent = text;
                    
                    messageDiv.appendChild(bubbleDiv);
                    
                    // Add sources if available
                    if (sources && sources.length > 0) {
                        const sourcesDiv = document.createElement('div');
                        sourcesDiv.className = 'message-sources';
                        sourcesDiv.textContent = `Sources: ${sources.join(', ')}`;
                        messageDiv.appendChild(sourcesDiv);
                    }
                    
                    this.messagesContainer.appendChild(messageDiv);
                    this.messagesContainer.scrollTop = this.messagesContainer.scrollHeight;
                }
                
                setInputDisabled(disabled) {
                    this.messageInput.disabled = disabled;
                    this.sendButton.disabled = disabled;
                }
                
                showTypingIndicator() {
                    this.typingIndicator.style.display = 'block';
                    this.messagesContainer.scrollTop = this.messagesContainer.scrollHeight;
                }
                
                hideTypingIndicator() {
                    this.typingIndicator.style.display = 'none';
                }
            }
            
            // Initialize chatbot
            new Chatbot();
        </script>
    </body>
    </html>
    """
    
    # Return HTML response (ngrok bypass headers added by middleware)
    return HTMLResponse(content=html_content)

@app.get("/", response_class=HTMLResponse)
async def admin_interface():
    """Admin interface for document management"""
    return """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>FutureClarity Chatbot Admin</title>
        <style>
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                margin: 0;
                padding: 20px;
                background: #f5f7fa;
            }
            .container {
                max-width: 1200px;
                margin: 0 auto;
                background: white;
                border-radius: 12px;
                padding: 30px;
                box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
            }
            h1 {
                color: #333;
                margin-bottom: 30px;
                text-align: center;
            }
            .section {
                margin-bottom: 40px;
            }
            .section h2 {
                color: #555;
                margin-bottom: 20px;
                padding-bottom: 10px;
                border-bottom: 2px solid #e1e5e9;
            }
            .upload-area {
                border: 2px dashed #007bff;
                border-radius: 8px;
                padding: 40px;
                text-align: center;
                margin-bottom: 20px;
                transition: background 0.2s;
            }
            .upload-area:hover {
                background: #f8f9ff;
            }
            .upload-area input[type="file"] {
                display: none;
            }
            .upload-btn {
                background: #007bff;
                color: white;
                padding: 12px 24px;
                border: none;
                border-radius: 6px;
                cursor: pointer;
                font-size: 16px;
                transition: background 0.2s;
            }
            .upload-btn:hover {
                background: #0056b3;
            }
            .documents-list {
                background: #f8f9fa;
                border-radius: 8px;
                padding: 20px;
            }
            .document-item {
                background: white;
                border-radius: 6px;
                padding: 15px;
                margin-bottom: 10px;
                display: flex;
                justify-content: space-between;
                align-items: center;
                border: 1px solid #e1e5e9;
            }
            .document-info {
                flex: 1;
            }
            .document-name {
                font-weight: 600;
                color: #333;
            }
            .document-meta {
                font-size: 0.9em;
                color: #666;
                margin-top: 4px;
            }
            .delete-btn {
                background: #dc3545;
                color: white;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
                cursor: pointer;
                font-size: 14px;
            }
            .delete-btn:hover {
                background: #c82333;
            }
            .chatbot-preview {
                border: 1px solid #e1e5e9;
                border-radius: 8px;
                overflow: hidden;
                margin-top: 20px;
            }
            .embed-code {
                background: #f8f9fa;
                border: 1px solid #e1e5e9;
                border-radius: 6px;
                padding: 15px;
                font-family: monospace;
                font-size: 14px;
                color: #333;
                overflow-x: auto;
                margin-top: 10px;
            }
            .status {
                padding: 10px;
                border-radius: 6px;
                margin-bottom: 20px;
                text-align: center;
                font-weight: 500;
            }
            .status.success {
                background: #d4edda;
                color: #155724;
                border: 1px solid #c3e6cb;
            }
            .status.error {
                background: #f8d7da;
                color: #721c24;
                border: 1px solid #f5c6cb;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>FutureClarity Chatbot Admin</h1>
            
            <div id="statusMessage"></div>
            
            <div class="section">
                <h2>Upload Documents</h2>
                <div class="upload-area" onclick="document.getElementById('fileInput').click()">
                    <p>Click to upload documents</p>
                    <p style="color: #666; font-size: 0.9em;">Supported formats: PDF, DOCX, TXT, MD</p>
                    <input type="file" id="fileInput" accept=".pdf,.docx,.txt,.md" />
                    <button type="button" class="upload-btn">Choose Files</button>
                </div>
            </div>
            
            <div class="section">
                <h2>Document Library</h2>
                <div class="documents-list" id="documentsList">
                    <p>Loading documents...</p>
                </div>
            </div>
            
            <div class="section">
                <h2>Chatbot Preview</h2>
                <iframe src="/embed" width="100%" height="600" class="chatbot-preview"></iframe>
            </div>
            
            <div class="section">
                <h2>Embed Code</h2>
                <p>Copy this code to embed the chatbot on your website:</p>
                <div class="embed-code" id="embedCode">
                    &lt;iframe src="http://localhost:8000/embed" width="400" height="600" frameborder="0"&gt;&lt;/iframe&gt;
                </div>
            </div>
        </div>
        
        <script>
            // Load documents on page load
            document.addEventListener('DOMContentLoaded', function() {
                loadDocuments();
                
                // Set up file upload
                document.getElementById('fileInput').addEventListener('change', function(e) {
                    if (e.target.files.length > 0) {
                        uploadFile(e.target.files[0]);
                    }
                });
                
                // Update embed code with current host
                const embedCode = document.getElementById('embedCode');
                embedCode.innerHTML = `&lt;iframe src="${window.location.origin}/embed" width="400" height="600" frameborder="0"&gt;&lt;/iframe&gt;`;
            });
            
            async function loadDocuments() {
                try {
                    const response = await fetch('/api/documents');
                    const documents = await response.json();
                    
                    const documentsList = document.getElementById('documentsList');
                    
                    if (Object.keys(documents).length === 0) {
                        documentsList.innerHTML = '<p>No documents uploaded yet.</p>';
                        return;
                    }
                    
                    documentsList.innerHTML = '';
                    
                    for (const [filename, info] of Object.entries(documents)) {
                        const docItem = document.createElement('div');
                        docItem.className = 'document-item';
                        
                        docItem.innerHTML = `
                            <div class="document-info">
                                <div class="document-name">${filename}</div>
                                <div class="document-meta">
                                    ${Math.round(info.size / 1024)} KB • ${info.chunks} chunks • 
                                    Processed: ${new Date(info.processed_at).toLocaleString()}
                                </div>
                            </div>
                            <button class="delete-btn" onclick="deleteDocument('${filename}')">Delete</button>
                        `;
                        
                        documentsList.appendChild(docItem);
                    }
                } catch (error) {
                    console.error('Error loading documents:', error);
                    showStatus('Error loading documents', 'error');
                }
            }
            
            async function uploadFile(file) {
                const formData = new FormData();
                formData.append('file', file);
                
                try {
                    showStatus('Uploading and processing document...', 'success');
                    
                    const response = await fetch('/api/upload', {
                        method: 'POST',
                        body: formData
                    });
                    
                    const result = await response.json();
                    
                    if (response.ok) {
                        showStatus('Document uploaded and processed successfully!', 'success');
                        loadDocuments();
                    } else {
                        showStatus('Error: ' + result.detail, 'error');
                    }
                } catch (error) {
                    console.error('Upload error:', error);
                    showStatus('Error uploading document', 'error');
                }
                
                // Reset file input
                document.getElementById('fileInput').value = '';
            }
            
            async function deleteDocument(filename) {
                if (!confirm(`Are you sure you want to delete "${filename}"?`)) {
                    return;
                }
                
                try {
                    const response = await fetch(`/api/documents/${encodeURIComponent(filename)}`, {
                        method: 'DELETE'
                    });
                    
                    if (response.ok) {
                        showStatus('Document deleted successfully', 'success');
                        loadDocuments();
                    } else {
                        showStatus('Error deleting document', 'error');
                    }
                } catch (error) {
                    console.error('Delete error:', error);
                    showStatus('Error deleting document', 'error');
                }
            }
            
            function showStatus(message, type) {
                const statusDiv = document.getElementById('statusMessage');
                statusDiv.innerHTML = `<div class="status ${type}">${message}</div>`;
                
                setTimeout(() => {
                    statusDiv.innerHTML = '';
                }, 5000);
            }
        </script>
    </body>
    </html>
    """

if __name__ == "__main__":
    print(f"""
    🚀 FutureClarity Chatbot System Starting...
    
    📋 Setup Instructions:
    1. Install Ollama: https://ollama.ai/download
    2. Pull a model: ollama pull mistral:7b
    3. Install Python dependencies: pip install -r requirements.txt
    4. Run this script: python chatbot_app.py
    
    🌐 Access Points:
    - Admin Interface: http://localhost:3000
    - Chatbot Embed: http://localhost:3000/embed
    - API Documentation: http://localhost:3000/docs
    
    📚 Document Management:
    - Upload documents via the admin interface
    - Supported formats: PDF, DOCX, TXT, MD
    - Documents are automatically processed and indexed
    
    🎯 Embedding on Websites:
    Use this iframe code:
    <iframe src="http://localhost:3000/embed" width="400" height="600" frameborder="0"></iframe>
    
    """)
    
    uvicorn.run(
        app,
        host=CONFIG["host"],
        port=CONFIG["port"],
        log_level="info"
    ) 